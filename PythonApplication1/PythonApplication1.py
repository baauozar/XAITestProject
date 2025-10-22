
from flask import Flask, request, jsonify
from flask_cors import CORS
from sentence_transformers import SentenceTransformer, util
import re
from pdfminer.high_level import extract_text as pdf_extract
from docx import Document
import tempfile

app = Flask(__name__); CORS(app)
model = SentenceTransformer("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()

def sbert_score(cv_text: str, job_text: str) -> float:
    a, b = normalize(cv_text), normalize(job_text)
    emb = model.encode([a, b], convert_to_tensor=True, normalize_embeddings=True)
    cos = float(util.cos_sim(emb[0], emb[1]))
    return (cos + 1.0) * 50.0

def explain_terms(cv_text: str, job_text: str, top_k: int = 10):
    tok = r"[a-zA-ZöçşığüÖÇŞİĞÜ0-9\+\#\-]{2,}"
    def clean_terms(s):
        import re as _re
        terms = _re.findall(tok, s.lower())
        return { _re.sub(r"^[^\w]+|[^\w]+$", "", t) for t in terms if t }
    a = clean_terms(cv_text); b = clean_terms(job_text)
    inter = list(a.intersection(b)); inter.sort(key=len, reverse=True)
    return inter[:top_k]

@app.get("/health")
def health(): return jsonify(status="ok")

@app.post("/score")
def post_score():
    data = request.get_json(force=True) or {}
    cv, job = data.get("cv_text", ""), data.get("job_text", "")
    base = round(sbert_score(cv, job), 2)
    final = max(0.0, min(100.0, base))
    reasons = []
    terms = explain_terms(cv, job)
    if terms: reasons.append(f"Ortak terimler / Overlaps: {', '.join(terms)}")
    reasons.append(f"Temel SBERT skor: {base}")
    return jsonify({"score": final, "base": base, "adjustment": 0, "explanation": reasons})

@app.post("/extract")
def extract():
    f = request.files.get("file")
    if not f:
        return jsonify(error="file is required"), 400
    filename = f.filename or ""
    content_type = f.mimetype or ""
    try:
        if filename.lower().endswith(".txt") or content_type.startswith("text/"):
            text = f.read().decode("utf-8", errors="ignore")
            return jsonify(text=text, format="txt")
        elif filename.lower().endswith(".pdf") or content_type == "application/pdf":
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
                f.save(tmp.name)
                text = pdf_extract(tmp.name) or ""
            return jsonify(text=text, format="pdf")
        elif filename.lower().endswith(".docx") or content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                f.save(tmp.name)
                doc = Document(tmp.name)
                parts = []
                for p in doc.paragraphs:
                    parts.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        parts.append(" ".join(cell.text for cell in row.cells))
                text = "\n".join(parts)
            return jsonify(text=text, format="docx")
        else:
            return jsonify(error="unsupported file type"), 415
    except Exception as e:
        return jsonify(error=str(e)), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8001)
